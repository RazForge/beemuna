import hashlib
import html
import io
import re
import uuid
from pathlib import Path

from docx import Document as DocxDocument
import pymupdf
from fastapi import HTTPException
from markdown_it import MarkdownIt

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".docx": "docx",
    ".html": "text",
    ".htm": "text",
    ".csv": "text",
    ".json": "text",
}

CHUNK_SIZE = 1200
CHUNK_OVERLAP = 200


def compute_checksum(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def extract_text(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = []
        with pymupdf.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                text.append(page.get_text())
        return "\n".join(text)
    if ext == ".docx":
        doc = DocxDocument(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text or "" for cell in row.cells))
        return "\n".join(parts)
    if ext in (".md", ".markdown"):
        text = data.decode("utf-8", errors="replace")
        return MarkdownIt("commonmark", {"html": True}).render(text)
    if ext == ".txt":
        return data.decode("utf-8", errors="replace")
    return data.decode("utf-8", errors="replace")


def _strip_html(text: str) -> str:
    text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.S | re.I)
    text = re.sub(r"<script[^>]*>.*?</script>", " ", text, flags=re.S | re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    return html.unescape(text)


def clean_text(raw: str) -> str:
    text = _strip_html(raw)
    text = re.sub(r"[\r\n\t]+", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = clean_text(text)
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            breakpoint = text.rfind(".", start, end)
            if breakpoint > start + chunk_size // 2:
                end = breakpoint + 1
            else:
                end = text.rfind(" ", start, end)
                if end <= start:
                    end = start + chunk_size
        chunks.append(text[start:end].strip())
        start = max(start + 1, end - overlap)
        if start >= len(text) - overlap:
            break
        if len(chunks) > 500:
            break
    return chunks


def rough_token_count(text: str) -> int:
    return max(1, len(text.split()))


def save_upload(user_id: uuid.UUID, space_id: uuid.UUID, source_id: uuid.UUID, filename: str, data: bytes, upload_dir: str) -> str:
    _mime_check(data, filename)
    base = Path(upload_dir)
    rel = Path(str(user_id)) / str(space_id) / f"{source_id}{Path(filename).suffix}"
    path = base / rel
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    return str(rel)


def _mime_check(data: bytes, filename: str) -> None:
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    header = data[:16]
    if ext == ".pdf":
        if not header.startswith(b"%PDF"):
            raise HTTPException(status_code=400, detail="File type mismatch")
    elif ext == ".docx":
        if not header.startswith(b"PK"):
            raise HTTPException(status_code=400, detail="File type mismatch")
    elif ext in (".txt", ".md", ".markdown"):
        if header and not _is_text(header):
            raise HTTPException(status_code=400, detail="File type mismatch")


def _is_text(header: bytes) -> bool:
    try:
        sample = header.decode("utf-8", errors="replace")
    except Exception:
        return False
    return all(ord(c) < 128 or c in "\n\r\t" for c in sample if c != "\x00")